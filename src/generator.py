"""内容生成模块 - 支持Ollama和本地模型"""

import os
import requests
import warnings
from typing import Dict, List, Optional, Generator

# 禁用警告
warnings.filterwarnings("ignore")

from .config import get_config
from .literature_manager import LiteratureManager


class NSFCGenerator:
    """国自然申请书内容生成器"""
    
    # 各模块的系统提示
    SECTION_PROMPTS = {
        "立项依据": """你正在撰写国自然申请书的"立项依据"部分。

要求：
1. 研究背景与意义：阐明研究的科学意义和应用价值
2. 国内外研究现状：客观评述研究进展，指出存在问题
3. 研究缺口：明确指出现有研究的不足
4. 本研究必要性：论证开展研究的必要性和可行性
5. 文献引用：使用[1]、[2]格式标注

确保内容专业、论证严谨、逻辑清晰。""",

        "研究内容": """你正在撰写国自然申请书的"研究内容"部分。

要求：
1. 明确核心研究问题
2. 分点阐述具体研究内容（3-5个方面）
3. 界定研究范围和边界
4. 确保内容紧扣研究目标
5. 体现系统性和可操作性

使用"（1）"、"（2）"格式组织层次。""",

        "研究方案": """你正在撰写国自然申请书的"研究方案"部分。

要求：
1. 研究方法：说明主要研究方法
2. 技术路线：描述清晰的技术路线
3. 实验步骤：分阶段说明具体步骤
4. 关键技术：指出难点及解决策略
5. 进度安排：合理安排研究进度
6. 可行性：论证方案的科学性""",

        "创新点": """你正在撰写国自然申请书的"创新点"部分。

要求：
1. 提炼2-4个创新点
2. 区分类型：理论创新、方法创新、应用创新
3. 具体说明创新之处
4. 与现有研究对比，说明独特性
5. 避免泛泛而谈

格式：创新点一：...""",

        "预期成果": """你正在撰写国自然申请书的"预期成果"部分。

要求：
1. 学术成果：论文数量和级别
2. 知识产权：专利、软著等
3. 人才培养：研究生培养
4. 应用价值：技术应用前景
5. 指标需具体可量化""",

        "研究基础": """你正在撰写国自然申请书的"研究基础"部分。

要求：
1. 前期研究积累
2. 团队构成和专长
3. 实验条件和平台
4. 合作资源
5. 论证具备完成研究的能力"""
    }
    
    def __init__(self, literature_manager: LiteratureManager = None, use_local: bool = False):
        self.config = get_config()
        self.literature_manager = literature_manager or LiteratureManager(lazy_init=True)
        self.use_local = use_local
        self.local_model = None
        self.ollama_host = self.config.ollama.host
        self.model_name = self.config.ollama.model_name
        
        if use_local:
            try:
                from .local_inference import get_local_model
                self.local_model = get_local_model()
            except ImportError:
                print("⚠️ 无法导入本地推理模块，将回退到 Ollama 模式")
                self.use_local = False
    
    def _call_ollama(
        self,
        prompt: str,
        system_prompt: str = "",
        stream: bool = False
    ):
        """调用Ollama API"""
        url = f"{self.ollama_host}/api/generate"
        
        payload = {
            "model": self.model_name,
            "prompt": prompt,
            "system": system_prompt,
            "stream": stream,
            "options": {
                "temperature": self.config.generation.temperature,
                "top_p": self.config.generation.top_p,
                "num_predict": self.config.generation.max_tokens,
                "repeat_penalty": self.config.generation.repeat_penalty
            }
        }
        
        if stream:
            return self._stream_response(url, payload)
        else:
            try:
                response = requests.post(url, json=payload, timeout=50000)
                response.raise_for_status()
                return response.json().get('response', '')
            except requests.exceptions.ConnectionError:
                raise Exception("无法连接到Ollama，请确保Ollama服务正在运行 (ollama serve)")
            except Exception as e:
                raise Exception(f"Ollama调用失败: {str(e)}")
    
    def _call_local(
        self,
        prompt: str,
        system_prompt: str = "",
        stream: bool = False
    ):
        """调用本地模型"""
        if self.local_model is None:
            raise Exception("本地模型未初始化")
            
        return self.local_model.generate(
            prompt=prompt,
            system_prompt=system_prompt,
            max_new_tokens=self.config.generation.max_tokens,
            temperature=self.config.generation.temperature,
            top_p=self.config.generation.top_p,
            stream=stream
        )
    
    def _stream_response(self, url: str, payload: dict) -> Generator[str, None, None]:
        """流式响应"""
        import json
        
        try:
            with requests.post(url, json=payload, stream=True, timeout=50000) as response:
                response.raise_for_status()
                for line in response.iter_lines():
                    if line:
                        data = json.loads(line)
                        if 'response' in data:
                            yield data['response']
                        if data.get('done', False):
                            break
        except Exception as e:
            yield f"\n\n❌ 生成出错: {str(e)}"
    
    def generate_section(
        self,
        section_type: str,
        research_topic: str,
        additional_info: str = "",
        use_literature: bool = True,
        stream: bool = False
    ) -> str:
        """
        生成申请书模块
        """
        if section_type not in self.SECTION_PROMPTS:
            raise ValueError(f"不支持的模块: {section_type}")
        
        system_prompt = self.SECTION_PROMPTS[section_type]
        
        # 构建提示词
        prompt_parts = [f"研究主题：{research_topic}"]
        
        if use_literature:
            try:
                context = self.literature_manager.build_context(research_topic)
                if context:
                    prompt_parts.append(f"\n{context}")
            except Exception as e:
                print(f"获取文献上下文失败: {e}")
        
        if additional_info:
            prompt_parts.append(f"\n补充信息：{additional_info}")
        
        # === 修正了这里：外部使用单引号，内部可以使用双引号 ===
        prompt_parts.append(f'\n请撰写"{section_type}"部分：')
        
        full_prompt = "\n".join(prompt_parts)
        
        if self.use_local:
            return self._call_local(full_prompt, system_prompt, stream)
        else:
            return self._call_ollama(full_prompt, system_prompt, stream)
    
    def generate_full_proposal(
        self,
        research_topic: str,
        use_literature: bool = True,
        sections: List[str] = None
    ) -> Dict[str, str]:
        """生成完整申请书"""
        if sections is None:
            sections = list(self.SECTION_PROMPTS.keys())
        
        results = {}
        
        for section in sections:
            print(f"正在生成: {section}...")
            try:
                content = self.generate_section(
                    section_type=section,
                    research_topic=research_topic,
                    use_literature=use_literature,
                    stream=False
                )
                results[section] = content
                print(f"✓ {section} ({len(content)} 字)")
            except Exception as e:
                results[section] = f"❌ 生成失败: {str(e)}"
                print(f"✗ {section}: {str(e)}")
        
        return results
    
    def refine_section(
        self,
        section_type: str,
        original_content: str,
        feedback: str
    ) -> str:
        """根据反馈修改内容"""
        system_prompt = f"""你是国自然申请书写作专家。
请根据修改意见优化"{section_type}"部分的内容。
保持原有结构的合理部分，针对性地改进不足之处。"""
        
        prompt = f"""原始内容：
{original_content}

修改意见：
{feedback}

请输出修改后的完整"{section_type}"内容："""
        
        if self.use_local:
            return self._call_local(prompt, system_prompt)
        else:
            return self._call_ollama(prompt, system_prompt)
    
    def add_literature(self, file_paths: List[str]) -> Dict[str, int]:
        """添加文献"""
        return self.literature_manager.add_files(file_paths)
    
    def get_literature_stats(self) -> Dict:
        """获取文献统计"""
        return self.literature_manager.get_stats()


class ProposalExporter:
    """申请书导出器"""
    
    SECTION_ORDER = ["立项依据", "研究内容", "研究方案", "创新点", "预期成果", "研究基础"]
    
    @classmethod
    def to_markdown(cls, sections: Dict[str, str], title: str = "") -> str:
        """导出为Markdown"""
        lines = [f"# {title or '国自然科学基金申请书'}\n"]
        
        for section in cls.SECTION_ORDER:
            if section in sections:
                lines.append(f"\n## {section}\n")
                lines.append(sections[section])
                lines.append("\n")
        
        return "\n".join(lines)
    
    @classmethod
    def to_docx(cls, sections: Dict[str, str], output_path: str, title: str = ""):
        """导出为Word文档"""
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        title_para = doc.add_heading(title or '国自然科学基金申请书', 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for section in cls.SECTION_ORDER:
            if section in sections:
                doc.add_heading(section, 1)
                
                content = sections[section]
                if content and not content.startswith("❌"):
                    paragraphs = content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text.strip())
                            para.paragraph_format.first_line_indent = Inches(0.3)
                            para.paragraph_format.line_spacing = 1.5
        
        doc.save(output_path)
        return output_path
    
    @classmethod
    def save(cls, sections: Dict[str, str], output_path: str, title: str = ""):
        """保存文件"""
        ext = os.path.splitext(output_path)[1].lower()
        
        if ext == '.md':
            content = cls.to_markdown(sections, title)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
        elif ext in ['.docx', '.doc']:
            cls.to_docx(sections, output_path, title)
        else:
            with open(output_path, 'w', encoding='utf-8') as f:
                for section in cls.SECTION_ORDER:
                    if section in sections:
                        f.write(f"\n{'='*50}\n{section}\n{'='*50}\n")
                        f.write(sections[section])
                        f.write("\n")
        
        print(f"✓ 已保存至: {output_path}")
        return output_path
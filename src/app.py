"""Web应用模块 - 支持模型切换和标书审阅"""

import os
import sys
import warnings
import tempfile

warnings.filterwarnings("ignore")
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["GRADIO_ANALYTICS_ENABLED"] = "False"

import gradio as gr
from typing import Tuple, Optional, List, Dict
import requests

from .config import get_config
from .generator import NSFCGenerator, ProposalExporter
from .literature_manager import LiteratureManager
from .proposal_reviewer import ProposalReviewer, ReviewResult


# 优化后的自定义CSS样式
CUSTOM_CSS = """
/* ============ 全局样式 ============ */
.gradio-container {
    max-width: 100% !important;
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
}

/* ============ 标题区域 ============ */
.title-section {
    text-align: center;
    padding: 40px 20px;
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    color: white;
    border-radius: 16px;
    margin-bottom: 30px;
    box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
    position: relative;
    overflow: hidden;
}

.title-section::before {
    content: '';
    position: absolute;
    top: -50%;
    right: -50%;
    width: 200%;
    height: 200%;
    background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%);
    animation: pulse 15s ease-in-out infinite;
}

@keyframes pulse {
    0%, 100% { transform: translate(0, 0); }
    50% { transform: translate(-10%, -10%); }
}

.title-section h1 {
    margin: 0 0 10px 0;
    font-size: 2.5em;
    font-weight: 700;
    position: relative;
    z-index: 1;
    text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
}

.title-section p {
    margin: 0;
    font-size: 1.1em;
    opacity: 0.95;
    position: relative;
    z-index: 1;
}

/* ============ 卡片样式 ============ */
.card {
    background: white;
    border-radius: 12px;
    padding: 24px;
    box-shadow: 0 4px 15px rgba(0, 0, 0, 0.08);
    transition: all 0.3s ease;
    border: 1px solid rgba(102, 126, 234, 0.1);
}

.card:hover {
    box-shadow: 0 8px 25px rgba(0, 0, 0, 0.12);
    transform: translateY(-2px);
}

/* ============ Tab样式优化 ============ */
.tab-nav button {
    font-size: 1.05em !important;
    padding: 12px 24px !important;
    border-radius: 8px 8px 0 0 !important;
    transition: all 0.3s ease !important;
}

.tab-nav button[aria-selected="true"] {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
    color: white !important;
    font-weight: 600 !important;
}

/* ============ 按钮样式 ============ */
.button-primary {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
    border: none !important;
    color: white !important;
    padding: 12px 28px !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    font-size: 1em !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4) !important;
}

.button-primary:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px rgba(102, 126, 234, 0.5) !important;
}

.button-secondary {
    background: white !important;
    border: 2px solid #667eea !important;
    color: #667eea !important;
    padding: 10px 24px !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    transition: all 0.3s ease !important;
}

.button-secondary:hover {
    background: #f0f4ff !important;
    transform: translateY(-2px) !important;
}

/* ============ 输入框样式 ============ */
.input-box textarea,
.input-box input {
    border-radius: 8px !important;
    border: 2px solid #e0e7ff !important;
    padding: 12px !important;
    font-size: 1em !important;
    transition: all 0.3s ease !important;
}

.input-box textarea:focus,
.input-box input:focus {
    border-color: #667eea !important;
    box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1) !important;
}

/* ============ 评分显示样式 ============ */
.review-score {
    font-size: 1.8em;
    font-weight: bold;
    padding: 20px;
    border-radius: 12px;
    text-align: center;
    margin: 16px 0;
    transition: all 0.3s ease;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
}

.score-high {
    background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
    color: #155724;
    border-left: 5px solid #28a745;
}

.score-medium {
    background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
    color: #856404;
    border-left: 5px solid #ffc107;
}

.score-low {
    background: linear-gradient(135deg, #f8d7da 0%, #f5c6cb 100%);
    color: #721c24;
    border-left: 5px solid #dc3545;
}

/* ============ 状态卡片 ============ */
.status-card {
    background: linear-gradient(135deg, #f8f9ff 0%, #e8edff 100%);
    border-radius: 12px;
    padding: 20px;
    border-left: 4px solid #667eea;
    margin: 12px 0;
}

.status-card h3 {
    color: #667eea;
    margin-top: 0;
}

/* ============ 文件上传区域 ============ */
.file-upload-area {
    border: 2px dashed #667eea !important;
    border-radius: 12px !important;
    background: linear-gradient(135deg, #f8f9ff 0%, #ffffff 100%) !important;
    padding: 30px !important;
    transition: all 0.3s ease !important;
}

.file-upload-area:hover {
    border-color: #764ba2 !important;
    background: linear-gradient(135deg, #f0f4ff 0%, #faf5ff 100%) !important;
}

/* ============ 进度条 ============ */
.progress-bar {
    background: linear-gradient(90deg, #667eea 0%, #764ba2 100%) !important;
    border-radius: 10px !important;
    height: 8px !important;
}

/* ============ 结果文本框 ============ */
.output-box {
    background: #fafbfc !important;
    border-radius: 8px !important;
    border: 1px solid #e1e4e8 !important;
    font-family: 'Monaco', 'Menlo', 'Consolas', monospace !important;
    line-height: 1.6 !important;
}

/* ============ Accordion优化 ============ */
.accordion {
    border-radius: 8px !important;
    border: 1px solid #e0e7ff !important;
}

.accordion-header {
    background: #f8f9ff !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
}

/* ============ 标签样式 ============ */
.label-text {
    font-weight: 600 !important;
    color: #333 !important;
    font-size: 1.05em !important;
    margin-bottom: 8px !important;
}

/* ============ 分隔线 ============ */
.divider {
    height: 2px;
    background: linear-gradient(90deg, transparent 0%, #667eea 50%, transparent 100%);
    margin: 30px 0;
    border: none;
}

/* ============ 信息框 ============ */
.info-box {
    background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%);
    border-left: 4px solid #2196f3;
    border-radius: 8px;
    padding: 16px;
    margin: 16px 0;
}

.success-box {
    background: linear-gradient(135deg, #e8f5e9 0%, #c8e6c9 100%);
    border-left: 4px solid #4caf50;
    border-radius: 8px;
    padding: 16px;
    margin: 16px 0;
}

.warning-box {
    background: linear-gradient(135deg, #fff8e1 0%, #ffecb3 100%);
    border-left: 4px solid #ff9800;
    border-radius: 8px;
    padding: 16px;
    margin: 16px 0;
}

/* ============ 页脚样式 ============ */
.footer {
    text-align: center;
    padding: 30px 20px;
    color: #666;
    background: linear-gradient(135deg, #f8f9ff 0%, #ffffff 100%);
    border-radius: 12px;
    margin-top: 30px;
    box-shadow: 0 -4px 15px rgba(0, 0, 0, 0.05);
}

.footer strong {
    color: #667eea;
}

/* ============ 响应式设计 ============ */
@media (max-width: 768px) {
    .title-section h1 {
        font-size: 1.8em;
    }
    
    .title-section p {
        font-size: 0.95em;
    }
    
    .card {
        padding: 16px;
    }
}

/* ============ 动画效果 ============ */
@keyframes fadeIn {
    from {
        opacity: 0;
        transform: translateY(10px);
    }
    to {
        opacity: 1;
        transform: translateY(0);
    }
}

.fade-in {
    animation: fadeIn 0.5s ease;
}

/* ============ 滚动条样式 ============ */
::-webkit-scrollbar {
    width: 10px;
    height: 10px;
}

::-webkit-scrollbar-track {
    background: #f1f1f1;
    border-radius: 5px;
}

::-webkit-scrollbar-thumb {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    border-radius: 5px;
}

::-webkit-scrollbar-thumb:hover {
    background: linear-gradient(135deg, #764ba2 0%, #667eea 100%);
}
"""


class WebApp:
    """Gradio Web应用"""
    
    def __init__(self):
        self.config = get_config()
        self.literature_manager = None
        self.generator = None
        self.reviewer = None
        self.current_sections = {}
        self._initialized = False
        self.current_model_type = "ollama"
        self.current_model_name = self.config.ollama.model_name
        
        # 审阅结果缓存
        self.review_results: Dict[str, ReviewResult] = {}
    
    def _ensure_initialized(self, use_local: bool = False):
        """初始化生成器"""
        need_reinit = (
            not self._initialized or
            (use_local and not getattr(self.generator, 'use_local', False)) or
            (not use_local and getattr(self.generator, 'use_local', True))
        )
        
        if need_reinit:
            self.literature_manager = LiteratureManager(lazy_init=True)
            self.generator = NSFCGenerator(
                self.literature_manager,
                use_local=use_local
            )
            self.reviewer = ProposalReviewer(generator=self.generator)
            self._initialized = True
    
    def get_available_ollama_models(self) -> List[str]:
        """获取Ollama可用模型列表"""
        try:
            response = requests.get(
                f"{self.config.ollama.host}/api/tags",
                timeout=5
            )
            if response.status_code == 200:
                models = response.json().get('models', [])
                return [m['name'] for m in models]
        except Exception:
            pass
        return []
    
    def check_local_model(self) -> Tuple[bool, str]:
        """检查本地微调模型是否可用"""
        merged_path = self.config.paths.merged_model
        finetuned_path = self.config.paths.finetuned_model
        
        if os.path.exists(merged_path) and os.listdir(merged_path):
            return True, f"✅ 合并模型可用: {merged_path}"
        elif os.path.exists(finetuned_path) and os.listdir(finetuned_path):
            return True, f"✅ LoRA模型可用: {finetuned_path}"
        else:
            return False, "❌ 未找到微调模型"
    
    def check_ollama_status(self) -> Tuple[bool, str]:
        """检查Ollama状态"""
        try:
            response = requests.get(
                f"{self.config.ollama.host}/api/tags",
                timeout=5
            )
            if response.status_code == 200:
                models = [m['name'] for m in response.json().get('models', [])]
                if models:
                    return True, f"✅ Ollama运行中，可用模型: {', '.join(models)}"
                else:
                    return False, "⚠️ Ollama运行中，但没有模型"
        except Exception:
            pass
        return False, "❌ Ollama未运行"
    
    def get_model_status(self) -> str:
        """获取模型状态信息"""
        lines = ["<div class='status-card'>"]
        lines.append("<h3>📊 模型状态</h3>")
        
        ollama_ok, ollama_msg = self.check_ollama_status()
        status_icon = "🟢" if ollama_ok else "🔴"
        lines.append(f"<p>{status_icon} <strong>Ollama:</strong> {ollama_msg}</p>")
        
        local_ok, local_msg = self.check_local_model()
        status_icon = "🟢" if local_ok else "🔴"
        lines.append(f"<p>{status_icon} <strong>本地微调模型:</strong> {local_msg}</p>")
        
        lines.append(f"<p>🎯 <strong>当前使用:</strong> {self.current_model_type} - {self.current_model_name}</p>")
        lines.append("</div>")
        
        return "\n".join(lines)
    
    def switch_model(self, model_type: str, ollama_model: str) -> str:
        """切换模型"""
        if model_type == "本地微调模型":
            local_ok, local_msg = self.check_local_model()
            if not local_ok:
                return f"<div class='warning-box'>❌ 切换失败：{local_msg}</div>"
            
            self.current_model_type = "local"
            self.current_model_name = "微调模型"
            self._ensure_initialized(use_local=True)
            return f"<div class='success-box'>✅ 已切换到本地微调模型</div>"
        
        else:
            ollama_ok, ollama_msg = self.check_ollama_status()
            if not ollama_ok:
                return f"<div class='warning-box'>❌ 切换失败：{ollama_msg}</div>"
            
            available = self.get_available_ollama_models()
            if ollama_model not in available:
                return f"<div class='warning-box'>❌ 模型 {ollama_model} 不可用</div>"
            
            self.current_model_type = "ollama"
            self.current_model_name = ollama_model
            self.config.ollama.model_name = ollama_model
            self._ensure_initialized(use_local=False)
            return f"<div class='success-box'>✅ 已切换到 Ollama 模型: {ollama_model}</div>"
    
    def refresh_ollama_models(self):
        """刷新Ollama模型列表"""
        models = self.get_available_ollama_models()
        if models:
            return gr.Dropdown(choices=models, value=models[0])
        else:
            return gr.Dropdown(choices=["无可用模型"], value="无可用模型")
    
    def upload_literature(self, files) -> str:
        """上传文献"""
        if not files:
            return "<div class='warning-box'>⚠️ 请选择要上传的文件</div>"
        
        self._ensure_initialized(use_local=(self.current_model_type == "local"))
        
        try:
            file_paths = [f.name for f in files]
            results = self.generator.add_literature(file_paths)
            
            success = sum(1 for v in results.values() if v > 0)
            chunks = sum(results.values())
            stats = self.generator.get_literature_stats()
            
            return f"""<div class='success-box'>
<h3>📊 上传完成</h3>
<p>📁 <strong>处理结果:</strong> {success}/{len(files)} 个文件成功</p>
<p>📝 <strong>新增文本块:</strong> {chunks}</p>
<p>📚 <strong>文献库总计:</strong> {stats['total_documents']} 篇文献</p>
</div>"""
        except Exception as e:
            return f"<div class='warning-box'>❌ 上传失败: {str(e)}</div>"
    
    def generate_section(
        self,
        section_type: str,
        research_topic: str,
        additional_info: str,
        use_literature: bool,
        temperature: float
    ) -> str:
        """生成模块内容"""
        if not research_topic or not research_topic.strip():
            return "⚠️ 请输入研究主题"
        
        use_local = (self.current_model_type == "local")
        self._ensure_initialized(use_local=use_local)
        
        try:
            self.config.generation.temperature = temperature
            
            content = self.generator.generate_section(
                section_type=section_type,
                research_topic=research_topic.strip(),
                additional_info=additional_info.strip() if additional_info else "",
                use_literature=use_literature,
                stream=False
            )
            self.current_sections[section_type] = content
            return content
        except Exception as e:
            return f"❌ 生成失败: {str(e)}"
    
    def refine_content(
        self,
        section_type: str,
        original: str,
        feedback: str
    ) -> str:
        """修改内容"""
        if not original or not original.strip():
            return "⚠️ 请先生成内容"
        if not feedback or not feedback.strip():
            return "⚠️ 请输入修改意见"
        
        use_local = (self.current_model_type == "local")
        self._ensure_initialized(use_local=use_local)
        
        try:
            refined = self.generator.refine_section(
                section_type,
                original.strip(),
                feedback.strip()
            )
            self.current_sections[section_type] = refined
            return refined
        except Exception as e:
            return f"❌ 修改失败: {str(e)}"
    
    def generate_all_sections(
        self,
        research_topic: str,
        use_literature: bool,
        progress=gr.Progress()
    ) -> Tuple[str, str, str, str, str, str]:
        """生成所有模块"""
        if not research_topic or not research_topic.strip():
            msg = "⚠️ 请输入研究主题"
            return msg, msg, msg, msg, msg, msg
        
        use_local = (self.current_model_type == "local")
        self._ensure_initialized(use_local=use_local)
        
        sections_order = ["立项依据", "研究内容", "研究方案", "创新点", "预期成果", "研究基础"]
        results = []
        
        for i, section in enumerate(sections_order):
            progress((i + 1) / len(sections_order), desc=f"正在生成: {section}")
            try:
                content = self.generator.generate_section(
                    section_type=section,
                    research_topic=research_topic.strip(),
                    use_literature=use_literature,
                    stream=False
                )
                results.append(content)
                self.current_sections[section] = content
            except Exception as e:
                results.append(f"❌ 生成失败: {str(e)}")
        
        return tuple(results)
    
    # ========== 标书审阅功能 ==========
    
    def review_proposal(
        self,
        file,
        progress=gr.Progress()
    ) -> Tuple[str, str, str, str]:
        """审阅上传的标书"""
        if file is None:
            return "<div class='warning-box'>⚠️ 请上传标书文件</div>", "", "", ""
        
        use_local = (self.current_model_type == "local")
        self._ensure_initialized(use_local=use_local)
        
        try:
            file_path = file.name
            
            # 解析标书
            progress(0.1, desc="正在解析文档...")
            sections = self.reviewer.parse_proposal(file_path)
            
            if not sections:
                return "<div class='warning-box'>❌ 无法解析文档内容</div>", "", "", ""
            
            parsed_info = f"<div class='success-box'><h3>📄 文档解析成功</h3>"
            parsed_info += f"<p>识别到 <strong>{len(sections)}</strong> 个模块：</p><ul>"
            for name in sections.keys():
                parsed_info += f"<li>{name}</li>"
            parsed_info += "</ul></div>"
            
            # 审阅各模块
            self.review_results = {}
            section_list = list(sections.items())
            
            for i, (section_name, content) in enumerate(section_list):
                progress((i + 1) / len(section_list) * 0.8 + 0.1, desc=f"正在审阅: {section_name}")
                result = self.reviewer.review_section(section_name, content, use_model=True)
                self.review_results[section_name] = result
            
            # 生成报告
            progress(0.95, desc="正在生成报告...")
            report = self.reviewer.generate_review_report(self.review_results)
            revised = self.reviewer.generate_revised_proposal(self.review_results)
            
            # 计算总分
            avg_score = sum(r.score for r in self.review_results.values()) / len(self.review_results)
            
            # 根据分数选择样式
            if avg_score >= 8:
                score_class = "score-high"
            elif avg_score >= 6:
                score_class = "score-medium"
            else:
                score_class = "score-low"
            
            score_text = f"<div class='review-score {score_class}'>📊 综合评分: {avg_score:.1f}/10</div>"
            
            return parsed_info, score_text, report, revised
            
        except Exception as e:
            return f"<div class='warning-box'>❌ 审阅失败: {str(e)}</div>", "", "", ""
    
    def get_section_review(self, section_name: str) -> Tuple[str, str, str, str]:
        """获取单个模块的审阅详情"""
        if not self.review_results or section_name not in self.review_results:
            return "请先上传并审阅标书", "", "", ""
        
        result = self.review_results[section_name]
        
        score_text = f"评分: {result.score}/10"
        issues_text = "\n".join([f"• {issue}" for issue in result.issues])
        suggestions_text = "\n".join([f"• {s}" for s in result.suggestions])
        
        return score_text, issues_text, suggestions_text, result.revised_content
    
    def export_review_report(self) -> Optional[str]:
        """导出审阅报告"""
        if not self.review_results:
            return None
        
        try:
            report = self.reviewer.generate_review_report(self.review_results)
            
            fd, path = tempfile.mkstemp(suffix='.md')
            with os.fdopen(fd, 'w', encoding='utf-8') as f:
                f.write(report)
            return path
        except Exception as e:
            print(f"导出失败: {e}")
            return None
    
    def export_revised_proposal(self) -> Optional[str]:
        """导出修改后的标书"""
        if not self.review_results:
            return None
        
        try:
            revised = self.reviewer.generate_revised_proposal(self.review_results)
            
            fd, path = tempfile.mkstemp(suffix='.md')
            with os.fdopen(fd, 'w', encoding='utf-8') as f:
                f.write(revised)
            return path
        except Exception as e:
            print(f"导出失败: {e}")
            return None
    
    def export_revised_docx(self) -> Optional[str]:
        """导出修改后的标书为Word"""
        if not self.review_results:
            return None
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            title = doc.add_heading('国自然科学基金申请书（修改版）', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            section_order = ["立项依据", "研究内容", "研究方案", "创新点", "预期成果", "研究基础"]
            
            for section_name in section_order:
                if section_name in self.review_results:
                    result = self.review_results[section_name]
                    
                    doc.add_heading(section_name, 1)
                    
                    paragraphs = result.revised_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text.strip())
                            para.paragraph_format.first_line_indent = Inches(0.3)
                            para.paragraph_format.line_spacing = 1.5
            
            fd, path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            doc.save(path)
            return path
        except Exception as e:
            print(f"导出失败: {e}")
            return None
    
    def export_markdown(self, topic: str) -> Optional[str]:
        """导出Markdown"""
        if not self.current_sections:
            return None
        
        try:
            content = ProposalExporter.to_markdown(
                self.current_sections,
                topic if topic else "国自然申请书"
            )
            
            fd, path = tempfile.mkstemp(suffix='.md')
            with os.fdopen(fd, 'w', encoding='utf-8') as f:
                f.write(content)
            return path
        except Exception as e:
            print(f"导出失败: {e}")
            return None
    
    def export_word(self, topic: str) -> Optional[str]:
        """导出Word"""
        if not self.current_sections:
            return None
        
        try:
            fd, path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            
            ProposalExporter.to_docx(
                self.current_sections,
                path,
                topic if topic else "国自然申请书"
            )
            return path
        except Exception as e:
            print(f"导出失败: {e}")
            return None
    
    def build_interface(self) -> gr.Blocks:
        """构建界面"""
        
        initial_models = self.get_available_ollama_models()
        if not initial_models:
            initial_models = ["qwen2.5:7b", "qwen2.5:14b"]
        
        with gr.Blocks(
            title="国自然写作助手",
            theme=gr.themes.Soft(
                primary_hue="indigo",
                secondary_hue="purple",
                font=["Segoe UI", "sans-serif"]
            ),
            css=CUSTOM_CSS
        ) as demo:
            
            gr.HTML("""
            <div class="title-section fade-in">
                <h1>🎓 国自然科学基金申请书写作助手</h1>
                <p>✨ 智能写作 | 📋 标书审阅 | 🤖 支持本地微调模型和Ollama模型</p>
            </div>
            """)
            
            with gr.Tabs():
                # ========== Tab 0: 模型设置 ==========
                with gr.Tab("⚙️ 模型设置"):
                    gr.HTML("<div class='card fade-in'><h2 style='margin-top:0;color:#667eea;'>🔧 选择推理模型</h2></div>")
                    
                    with gr.Row():
                        with gr.Column(scale=1):
                            gr.HTML("<div class='card'>")
                            model_type = gr.Radio(
                                choices=["Ollama模型", "本地微调模型"],
                                value="Ollama模型",
                                label="🎯 模型类型",
                                elem_classes=["input-box"]
                            )
                            
                            gr.Markdown("---")
                            gr.Markdown("**📦 Ollama模型选择：**")
                            with gr.Row():
                                ollama_model = gr.Dropdown(
                                    choices=initial_models,
                                    value=initial_models[0] if initial_models else "qwen2.5:7b",
                                    label="选择模型",
                                    scale=3,
                                    elem_classes=["input-box"]
                                )
                                refresh_btn = gr.Button("🔄", scale=1, size="sm")
                            
                            gr.Markdown("---")
                            switch_btn = gr.Button(
                                "✅ 应用设置", 
                                variant="primary", 
                                size="lg",
                                elem_classes=["button-primary"]
                            )
                            switch_result = gr.HTML(label="结果")
                            gr.HTML("</div>")
                        
                        with gr.Column(scale=1):
                            status_display = gr.HTML(self.get_model_status())
                            refresh_status_btn = gr.Button(
                                "🔄 刷新状态",
                                elem_classes=["button-secondary"]
                            )
                    
                    refresh_btn.click(fn=self.refresh_ollama_models, outputs=[ollama_model])
                    switch_btn.click(fn=self.switch_model, inputs=[model_type, ollama_model], outputs=[switch_result])
                    refresh_status_btn.click(fn=self.get_model_status, outputs=[status_display])
                
                # ========== Tab 1: 标书审阅 ==========
                with gr.Tab("📋 标书审阅"):
                    gr.HTML("""
                    <div class='card fade-in'>
                        <h2 style='margin-top:0;color:#667eea;'>📤 上传您的申请书初稿</h2>
                        <p style='color:#666;'>上传Word或PDF格式的申请书，系统会自动分析各模块内容，给出专业的修改建议和优化后的版本。</p>
                    </div>
                    """)
                    
                    with gr.Row():
                        with gr.Column(scale=1):
                            gr.HTML("<div class='card'>")
                            proposal_file = gr.File(
                                label="📁 选择文件 (Word/PDF)",
                                file_types=[".docx", ".doc", ".pdf"],
                                height=180,
                                elem_classes=["file-upload-area"]
                            )
                            review_btn = gr.Button(
                                "🔍 开始审阅", 
                                variant="primary", 
                                size="lg",
                                elem_classes=["button-primary"]
                            )
                            
                            parse_result = gr.HTML(label="解析结果")
                            score_display = gr.HTML(label="综合评分")
                            gr.HTML("</div>")
                        
                        with gr.Column(scale=2):
                            gr.HTML("<div class='card'>")
                            with gr.Tabs():
                                with gr.Tab("📊 审阅报告"):
                                    review_report = gr.Markdown(
                                        label="",
                                        elem_classes=["output-box"]
                                    )
                                
                                with gr.Tab("📝 修改后版本"):
                                    revised_content = gr.Markdown(
                                        label="",
                                        elem_classes=["output-box"]
                                    )
                            gr.HTML("</div>")
                    
                    gr.HTML("<hr class='divider'>")
                    gr.HTML("<div class='card'><h3 style='color:#667eea;margin-top:0;'>📥 导出结果</h3>")
                    with gr.Row():
                        export_report_btn = gr.Button(
                            "📄 导出审阅报告", 
                            scale=1,
                            elem_classes=["button-secondary"]
                        )
                        export_revised_md_btn = gr.Button(
                            "📝 导出修改版(MD)", 
                            scale=1,
                            elem_classes=["button-secondary"]
                        )
                        export_revised_docx_btn = gr.Button(
                            "📄 导出修改版(Word)", 
                            scale=1,
                            elem_classes=["button-secondary"]
                        )
                        export_file = gr.File(label="📦 下载", scale=1)
                    gr.HTML("</div>")
                    
                    review_btn.click(
                        fn=self.review_proposal,
                        inputs=[proposal_file],
                        outputs=[parse_result, score_display, review_report, revised_content]
                    )
                    
                    export_report_btn.click(fn=self.export_review_report, outputs=[export_file])
                    export_revised_md_btn.click(fn=self.export_revised_proposal, outputs=[export_file])
                    export_revised_docx_btn.click(fn=self.export_revised_docx, outputs=[export_file])
                
                # ========== Tab 2: 文献管理 ==========
                with gr.Tab("📚 文献管理"):
                    gr.HTML("""
                    <div class='card fade-in'>
                        <h2 style='margin-top:0;color:#667eea;'>📤 上传参考文献</h2>
                        <p style='color:#666;'>支持PDF、Word、Markdown等多种格式，上传后系统会自动索引，在生成时智能引用</p>
                    </div>
                    """)
                    
                    with gr.Row():
                        with gr.Column(scale=1):
                            gr.HTML("<div class='card'>")
                            file_input = gr.File(
                                label="📂 拖拽文件到此处",
                                file_count="multiple",
                                file_types=[".pdf", ".docx", ".doc", ".md", ".txt"],
                                height=200,
                                elem_classes=["file-upload-area"]
                            )
                            upload_btn = gr.Button(
                                "📤 开始上传", 
                                variant="primary",
                                size="lg",
                                elem_classes=["button-primary"]
                            )
                            gr.HTML("</div>")
                        
                        with gr.Column(scale=1):
                            upload_output = gr.HTML(label="上传结果")
                    
                    upload_btn.click(fn=self.upload_literature, inputs=[file_input], outputs=[upload_output])
                
                # ========== Tab 3: 分模块生成 ==========
                with gr.Tab("✍️ 分模块生成"):
                    with gr.Row():
                        with gr.Column(scale=2):
                            gr.HTML("<div class='card'>")
                            topic_input = gr.Textbox(
                                label="🎯 研究主题 *",
                                placeholder="例如：基于深度学习的医学图像分析方法研究",
                                lines=2,
                                elem_classes=["input-box"]
                            )
                            
                            with gr.Row():
                                section_select = gr.Dropdown(
                                    label="📑 选择模块",
                                    choices=["立项依据", "研究内容", "研究方案",
                                            "创新点", "预期成果", "研究基础"],
                                    value="立项依据",
                                    scale=2,
                                    elem_classes=["input-box"]
                                )
                                use_lit_check = gr.Checkbox(
                                    label="📚 使用文献库", 
                                    value=True, 
                                    scale=1
                                )
                            
                            extra_input = gr.Textbox(
                                label="💡 补充信息（可选）",
                                placeholder="可以输入特定要求、关键技术等",
                                lines=2,
                                elem_classes=["input-box"]
                            )
                            
                            with gr.Accordion("🎚️ 高级参数", open=False):
                                temp_slider = gr.Slider(
                                    minimum=0.1, maximum=1.0, value=0.7, step=0.1,
                                    label="Temperature（创造性）",
                                    info="值越大生成越多样，越小越保守"
                                )
                            
                            with gr.Row():
                                gen_btn = gr.Button(
                                    "🚀 开始生成", 
                                    variant="primary", 
                                    scale=2,
                                    size="lg",
                                    elem_classes=["button-primary"]
                                )
                                clear_btn = gr.Button(
                                    "🗑️ 清空", 
                                    scale=1,
                                    elem_classes=["button-secondary"]
                                )
                            gr.HTML("</div>")
                        
                        with gr.Column(scale=3):
                            gr.HTML("<div class='card'>")
                            output_text = gr.Textbox(
                                label="📄 生成结果", 
                                lines=18, 
                                show_copy_button=True,
                                elem_classes=["output-box"]
                            )
                            
                            gr.Markdown("---")
                            with gr.Row():
                                feedback_input = gr.Textbox(
                                    label="✏️ 修改意见",
                                    placeholder="例如：增加更多技术细节、简化表述等",
                                    lines=1,
                                    scale=4,
                                    elem_classes=["input-box"]
                                )
                                refine_btn = gr.Button(
                                    "✨ 智能修改", 
                                    scale=1,
                                    elem_classes=["button-secondary"]
                                )
                            gr.HTML("</div>")
                    
                    gen_btn.click(
                        fn=self.generate_section,
                        inputs=[section_select, topic_input, extra_input, use_lit_check, temp_slider],
                        outputs=[output_text]
                    )
                    clear_btn.click(fn=lambda: "", outputs=[output_text])
                    refine_btn.click(
                        fn=self.refine_content,
                        inputs=[section_select, output_text, feedback_input],
                        outputs=[output_text]
                    )
                
                # ========== Tab 4: 完整生成 ==========
                with gr.Tab("📝 完整申请书"):
                    gr.HTML("""
                    <div class='card fade-in'>
                        <h2 style='margin-top:0;color:#667eea;'>🚀 一键生成完整申请书</h2>
                        <p style='color:#666;'>输入研究主题后，系统将自动生成所有必需模块的内容</p>
                    </div>
                    """)
                    
                    gr.HTML("<div class='card'>")
                    with gr.Row():
                        full_topic = gr.Textbox(
                            label="🎯 研究主题",
                            placeholder="请输入您的研究主题",
                            lines=1,
                            scale=4,
                            elem_classes=["input-box"]
                        )
                        full_use_lit = gr.Checkbox(
                            label="📚 使用文献库", 
                            value=True, 
                            scale=1
                        )
                        full_gen_btn = gr.Button(
                            "🚀 生成全部", 
                            variant="primary", 
                            scale=1,
                            size="lg",
                            elem_classes=["button-primary"]
                        )
                    gr.HTML("</div>")
                    
                    with gr.Row():
                        with gr.Column():
                            gr.HTML("<div class='card'>")
                            out_1 = gr.Textbox(
                                label="一、立项依据", 
                                lines=8, 
                                show_copy_button=True,
                                elem_classes=["output-box"]
                            )
                            out_2 = gr.Textbox(
                                label="二、研究内容", 
                                lines=6, 
                                show_copy_button=True,
                                elem_classes=["output-box"]
                            )
                            out_3 = gr.Textbox(
                                label="三、研究方案", 
                                lines=8, 
                                show_copy_button=True,
                                elem_classes=["output-box"]
                            )
                            gr.HTML("</div>")
                        with gr.Column():
                            gr.HTML("<div class='card'>")
                            out_4 = gr.Textbox(
                                label="四、创新点", 
                                lines=6, 
                                show_copy_button=True,
                                elem_classes=["output-box"]
                            )
                            out_5 = gr.Textbox(
                                label="五、预期成果", 
                                lines=5, 
                                show_copy_button=True,
                                elem_classes=["output-box"]
                            )
                            out_6 = gr.Textbox(
                                label="六、研究基础", 
                                lines=7, 
                                show_copy_button=True,
                                elem_classes=["output-box"]
                            )
                            gr.HTML("</div>")
                    
                    gr.HTML("<hr class='divider'>")
                    gr.HTML("<div class='card'><h3 style='color:#667eea;margin-top:0;'>📥 导出完整申请书</h3>")
                    with gr.Row():
                        exp_md_btn = gr.Button(
                            "📄 导出 Markdown",
                            elem_classes=["button-secondary"]
                        )
                        exp_docx_btn = gr.Button(
                            "📝 导出 Word",
                            elem_classes=["button-secondary"]
                        )
                        exp_file = gr.File(label="📦 下载")
                    gr.HTML("</div>")
                    
                    full_gen_btn.click(
                        fn=self.generate_all_sections,
                        inputs=[full_topic, full_use_lit],
                        outputs=[out_1, out_2, out_3, out_4, out_5, out_6]
                    )
                    exp_md_btn.click(fn=self.export_markdown, inputs=[full_topic], outputs=[exp_file])
                    exp_docx_btn.click(fn=self.export_word, inputs=[full_topic], outputs=[exp_file])
                
                # ========== Tab 5: 帮助 ==========
                with gr.Tab("❓ 帮助"):
                    gr.HTML("""
                    <div class='card fade-in'>
                        <h2 style='color:#667eea;margin-top:0;'>📖 功能说明</h2>
                        
                        <div class='info-box'>
                            <h3>📋 标书审阅</h3>
                            <ol>
                                <li>上传您的申请书初稿（Word或PDF格式）</li>
                                <li>系统自动识别各模块内容</li>
                                <li>AI给出专业评审意见和评分</li>
                                <li>自动生成修改后的版本</li>
                                <li>可导出审阅报告和修改版文档</li>
                            </ol>
                        </div>
                        
                        <div class='success-box'>
                            <h3>✍️ 智能写作</h3>
                            <ul>
                                <li><strong>分模块生成：</strong>逐个生成各个模块，可针对性优化</li>
                                <li><strong>完整生成：</strong>一键生成全部内容，快速出稿</li>
                                <li><strong>内容修改：</strong>根据您的意见智能优化内容</li>
                                <li><strong>文献支持：</strong>自动引用上传的参考文献</li>
                            </ul>
                        </div>
                        
                        <div class='warning-box'>
                            <h3>⚙️ 模型选择</h3>
                            <ul>
                                <li><strong>Ollama模型：</strong>使用Ollama管理的预训练模型（推荐Qwen2.5）</li>
                                <li><strong>本地微调模型：</strong>使用您自己微调后的专用模型</li>
                            </ul>
                            <p><em>💡 提示：首次使用前请先在"模型设置"中选择并应用模型</em></p>
                        </div>
                        
                        <div class='info-box'>
                            <h3>📚 文献管理</h3>
                            <p>上传参考文献后，系统会：</p>
                            <ul>
                                <li>自动提取文献内容并建立索引</li>
                                <li>在生成时智能检索相关内容</li>
                                <li>自然地融入到申请书中</li>
                            </ul>
                            <p><em>支持格式：PDF、Word、Markdown、TXT</em></p>
                        </div>
                        
                        <hr style='margin: 30px 0; border: none; border-top: 2px solid #e0e7ff;'>
                        
                        <h3 style='color:#667eea;'>🎯 使用技巧</h3>
                        <ul>
                            <li>📝 <strong>研究主题要具体：</strong>提供详细的研究方向和关键词</li>
                            <li>💡 <strong>补充信息要明确：</strong>说明特殊要求和技术细节</li>
                            <li>🔄 <strong>反复修改优化：</strong>利用修改功能逐步完善内容</li>
                            <li>📚 <strong>上传高质量文献：</strong>文献质量影响生成效果</li>
                            <li>⚡ <strong>调整Temperature：</strong>创新性内容用高值，规范性内容用低值</li>
                        </ul>
                    </div>
                    """)
            
            gr.HTML("""
            <div class="footer fade-in">
                <p style='font-size:1.1em;'>🎓 <strong>国自然写作助手 v1.1</strong></p>
                <p style='margin:5px 0;'>支持标书审阅 • 智能写作 • 模型切换</p>
                <p style='margin:10px 0;color:#999;font-size:0.9em;'>Powered by AI • Made with ❤️</p>
            </div>
            """)
        
        return demo
    
    def launch(self):
        """启动应用"""
        print("=" * 50)
        print("🚀 启动Web应用")
        print("=" * 50)
        
        demo = self.build_interface()
        
        demo.launch(
            server_name="127.0.0.1",
            server_port=self.config.webapp.port,
            share=False,
            inbrowser=True,
            show_error=True
        )


def run_webapp():
    """启动Web应用"""
    app = WebApp()
    app.launch()